"""Business objective: create equivalent full-invoice examples across formats from one source of truth.

Technical description: renders parties, header, line items, discounts/tax/shipping, and balance due into JSON, CSV, TXT, PDF, DOCX, and minimal XLSX inputs for format-invariance and persistence tests.
"""

from __future__ import annotations

import csv
import json
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen.canvas import Canvas

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "data" / "examples" / "input"
SOURCE = {
    "invoice_number": "19283746552",
    "invoice_date": "2025-07-01",
    "due_date": "2025-07-16",
    "payment_terms": "Due in 15 days",
    "currency": "USD",
    "parties": [
        {
            "role": "seller",
            "name": "Testinger GmbH",
            "address_lines": ["123 Test Street", "Testingen, Testern", "0815"],
            "phone": "+00 1234 5678 00",
            "email": "email@testinger.ts",
            "website": "testinger.ts",
        },
        {
            "role": "bill_to",
            "name": "Recipient Corp.",
            "contact_name": "Mr Obertestinger",
            "address_lines": ["123 Empfängergasse", "0815 Recipstan"],
            "phone": "+00 1234 8876 0912",
            "email": "recip@recipcopr.co",
        },
        {
            "role": "ship_to",
            "name": "Recipient Corp.",
            "contact_name": "Mr Obertestinger",
            "address_lines": ["123 Empfängergasse", "0815 Austria"],
            "phone": "+00 1234 8876 0912",
            "email": "recip@recipcopr.co",
        },
    ],
    "financials": {
        "subtotal": "280.00",
        "discount_total": "13.00",
        "subtotal_after_discount": "267.00",
        "tax_rate_percent": "20.00",
        "tax_total": "53.40",
        "shipping_total": "12.00",
        "amount_due": "332.40",
    },
    "lines": [
        {"description": "Highlife Steel Accessories", "quantity": 1, "unit_price": "10.00", "total": "10.00"},
        {"description": "Sneaker “Unstoppable”", "quantity": 2, "unit_price": "11.00", "total": "22.00"},
        {"description": "T-Shirt White “Polarbear”", "quantity": 3, "unit_price": "12.00", "total": "36.00"},
        {"description": "T-Shirt Beige “Grizzly”", "quantity": 4, "unit_price": "13.00", "total": "52.00"},
        {"description": "Shorts “El Camino”", "quantity": 5, "unit_price": "14.00", "total": "70.00"},
        {"description": "Socks, black", "quantity": 6, "unit_price": "15.00", "total": "90.00"},
    ],
}


def _party(role: str) -> dict[str, object]:
    return next(item for item in SOURCE["parties"] if item["role"] == role)


def _metadata_rows() -> list[list[object]]:
    seller = _party("seller")
    bill = _party("bill_to")
    ship = _party("ship_to")
    financials = SOURCE["financials"]
    return [
        ["INVOICE NUMBER", SOURCE["invoice_number"]],
        ["INVOICE DATE", SOURCE["invoice_date"]],
        ["DUE DATE", SOURCE["due_date"]],
        ["PAYMENT TERMS", SOURCE["payment_terms"]],
        ["CURRENCY", SOURCE["currency"]],
        ["SELLER NAME", seller["name"]],
        ["SELLER ADDRESS", "|".join(seller["address_lines"])],
        ["SELLER PHONE", seller["phone"]],
        ["SELLER EMAIL", seller["email"]],
        ["SELLER WEBSITE", seller["website"]],
        ["BILL TO NAME", bill["name"]],
        ["BILL TO CONTACT", bill["contact_name"]],
        ["BILL TO ADDRESS", "|".join(bill["address_lines"])],
        ["BILL TO PHONE", bill["phone"]],
        ["BILL TO EMAIL", bill["email"]],
        ["SHIP TO NAME", ship["name"]],
        ["SHIP TO CONTACT", ship["contact_name"]],
        ["SHIP TO ADDRESS", "|".join(ship["address_lines"])],
        ["SHIP TO PHONE", ship["phone"]],
        ["SHIP TO EMAIL", ship["email"]],
        [],
        ["DESCRIPTION", "QTY", "UNIT PRICE", "TOTAL"],
        *[[row["description"], row["quantity"], row["unit_price"], row["total"]] for row in SOURCE["lines"]],
        ["SUBTOTAL", "", "", financials["subtotal"]],
        ["DISCOUNT", "", "", financials["discount_total"]],
        ["SUBTOTAL AFTER DISCOUNT", "", "", financials["subtotal_after_discount"]],
        ["TAX RATE PERCENT", "", "", financials["tax_rate_percent"]],
        ["TAX TOTAL", "", "", financials["tax_total"]],
        ["SHIPPING TOTAL", "", "", financials["shipping_total"]],
        ["AMOUNT DUE", "", "", financials["amount_due"]],
    ]


def write_json() -> None:
    payload = {
        "business_objective": "Provide a deterministic full-invoice JSON fixture.",
        "technical_description": "The same parties, commercial totals, and six challenge invoice lines are rendered into every supported format.",
        "invoice": {
            "invoice_number": SOURCE["invoice_number"],
            "invoice_date": SOURCE["invoice_date"],
            "due_date": SOURCE["due_date"],
            "payment_terms": SOURCE["payment_terms"],
            "currency": SOURCE["currency"],
            "parties": SOURCE["parties"],
            "financials": {"currency": SOURCE["currency"], **SOURCE["financials"]},
        },
        "lines": SOURCE["lines"],
    }
    (OUTPUT / "equivalent_invoice.json").write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def write_csv_and_txt() -> None:
    rows = _metadata_rows()
    for name, delimiter in (("equivalent_invoice.csv", ","), ("equivalent_invoice.txt", "\t")):
        with (OUTPUT / name).open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle, delimiter=delimiter)
            writer.writerows(rows)


def _draw_party(canvas: Canvas, x: float, y: float, party: dict[str, object], *, heading: str | None = None) -> None:
    if heading:
        canvas.setFont("Helvetica-Bold", 9)
        canvas.drawString(x, y, heading)
        y -= 16
    canvas.setFont("Helvetica", 9)
    if party.get("contact_name"):
        canvas.drawString(x, y, str(party["contact_name"]))
        y -= 14
    canvas.drawString(x, y, str(party["name"]))
    y -= 14
    for address in party.get("address_lines", []):
        canvas.drawString(x, y, str(address))
        y -= 14
    if party.get("phone"):
        canvas.drawString(x, y, str(party["phone"]))
        y -= 14
    if party.get("email"):
        canvas.drawString(x, y, str(party["email"]))
        y -= 14
    if party.get("website"):
        canvas.drawString(x, y, str(party["website"]))


def write_pdf() -> None:
    path = OUTPUT / "equivalent_invoice.pdf"
    canvas = Canvas(str(path), pagesize=A4)
    width, height = A4
    canvas.setFont("Helvetica-Bold", 18)
    canvas.drawString(45, height - 50, "INVOICE")
    canvas.setFont("Helvetica", 9)
    canvas.drawRightString(width - 45, height - 50, "01. July 2025")
    canvas.drawRightString(width - 45, height - 66, str(SOURCE["invoice_number"]))
    canvas.drawRightString(width - 45, height - 84, str(SOURCE["payment_terms"]))

    _draw_party(canvas, 45, height - 125, _party("seller"))
    _draw_party(canvas, 235, height - 125, _party("bill_to"), heading="BILL TO")
    _draw_party(canvas, 405, height - 125, _party("ship_to"), heading="SHIP TO")

    y = height - 270
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawString(235, y, "DESCRIPTION")
    canvas.drawString(405, y, "QTY")
    canvas.drawString(455, y, "UNIT PRICE")
    canvas.drawString(535, y, "TOTAL")
    y -= 18
    canvas.setFont("Helvetica", 9)
    for row in SOURCE["lines"]:
        description = str(row["description"]).replace("“", '"').replace("”", '"')
        canvas.drawString(235, y, description)
        canvas.drawRightString(430, y, str(row["quantity"]))
        canvas.drawRightString(510, y, str(row["unit_price"]))
        canvas.drawRightString(570, y, str(row["total"]))
        y -= 18
    f = SOURCE["financials"]
    for label, value in (
        ("SUBTOTAL", f["subtotal"]),
        ("DISCOUNT", f["discount_total"]),
        ("SUBTOTAL AFTER DISCOUNT", f["subtotal_after_discount"]),
        ("TAX RATE", f'{f["tax_rate_percent"]}%'),
        ("TOTAL TAX", f["tax_total"]),
        ("SHIPPING/HANDLING", f["shipping_total"]),
        ("AMOUNT DUE", f'$ {f["amount_due"]}'),
    ):
        canvas.setFont("Helvetica-Bold" if label == "AMOUNT DUE" else "Helvetica", 9)
        canvas.drawRightString(500, y, label)
        canvas.drawRightString(570, y, str(value))
        y -= 16
    canvas.save()


def write_docx() -> None:
    document = Document()
    document.add_heading("INVOICE", level=1)
    metadata = document.add_table(rows=0, cols=2)
    for row in _metadata_rows()[:20]:
        if not row:
            continue
        cells = metadata.add_row().cells
        cells[0].text = str(row[0])
        cells[1].text = str(row[1])
    table = document.add_table(rows=1, cols=4)
    for index, value in enumerate(["DESCRIPTION", "QTY", "UNIT PRICE", "TOTAL"]):
        table.rows[0].cells[index].text = value
    for row in SOURCE["lines"]:
        cells = table.add_row().cells
        cells[0].text = str(row["description"])
        cells[1].text = str(row["quantity"])
        cells[2].text = str(row["unit_price"])
        cells[3].text = str(row["total"])
    f = SOURCE["financials"]
    for label, value in (
        ("SUBTOTAL", f["subtotal"]),
        ("DISCOUNT", f["discount_total"]),
        ("SUBTOTAL AFTER DISCOUNT", f["subtotal_after_discount"]),
        ("TAX RATE PERCENT", f["tax_rate_percent"]),
        ("TAX TOTAL", f["tax_total"]),
        ("SHIPPING TOTAL", f["shipping_total"]),
        ("AMOUNT DUE", f["amount_due"]),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[3].text = str(value)
    document.save(OUTPUT / "equivalent_invoice.docx")


def _column_name(index: int) -> str:
    value = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        value = chr(65 + remainder) + value
    return value


def write_xlsx() -> None:
    rows = _metadata_rows()
    sheet_rows: list[str] = []
    for row_index, row in enumerate(rows, start=1):
        cells: list[str] = []
        for column_index, value in enumerate(row, start=1):
            ref = f"{_column_name(column_index)}{row_index}"
            is_header = bool(row and str(row[0]) == "DESCRIPTION")
            style = 1 if is_header else (2 if column_index >= 3 and value != "" else 0)
            style_attr = f' s="{style}"' if style else ""
            text_value = str(value)
            try:
                numeric = float(text_value) if text_value and column_index >= 2 and str(row[0]) not in {
                    "INVOICE NUMBER", "SELLER PHONE", "BILL TO PHONE", "SHIP TO PHONE"
                } else None
            except ValueError:
                numeric = None
            if numeric is not None:
                cells.append(f'<c r="{ref}"{style_attr}><v>{numeric}</v></c>')
            else:
                cells.append(f'<c r="{ref}" t="inlineStr"{style_attr}><is><t>{escape(text_value)}</t></is></c>')
        sheet_rows.append(f'<row r="{row_index}" ht="20" customHeight="1">{"".join(cells)}</row>')
    worksheet = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        '<cols><col min="1" max="1" width="30" customWidth="1"/><col min="2" max="2" width="42" customWidth="1"/>'
        '<col min="3" max="3" width="14" customWidth="1"/><col min="4" max="4" width="14" customWidth="1"/></cols>'
        f'<sheetData>{"".join(sheet_rows)}</sheetData>'
        '</worksheet>'
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        '<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>'
        '</Types>'
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
        '</Relationships>'
    )
    workbook = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<sheets><sheet name="Invoice" sheetId="1" r:id="rId1"/></sheets></workbook>'
    )
    workbook_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
        '</Relationships>'
    )
    styles = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        '<numFmts count="1"><numFmt numFmtId="164" formatCode="0.00"/></numFmts>'
        '<fonts count="2"><font><sz val="11"/><name val="Calibri"/></font>'
        '<font><b/><color rgb="FFFFFFFF"/><sz val="11"/><name val="Calibri"/></font></fonts>'
        '<fills count="3"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill>'
        '<fill><patternFill patternType="solid"><fgColor rgb="FF1F5FBF"/><bgColor indexed="64"/></patternFill></fill></fills>'
        '<borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>'
        '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>'
        '<cellXfs count="3"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/>'
        '<xf numFmtId="0" fontId="1" fillId="2" borderId="0" xfId="0" applyFill="1" applyFont="1"/>'
        '<xf numFmtId="164" fontId="0" fillId="0" borderId="0" xfId="0" applyNumberFormat="1"/></cellXfs>'
        '<cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>'
        '</styleSheet>'
    )
    with zipfile.ZipFile(OUTPUT / "equivalent_invoice.xlsx", "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", root_rels)
        archive.writestr("xl/workbook.xml", workbook)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        archive.writestr("xl/worksheets/sheet1.xml", worksheet)
        archive.writestr("xl/styles.xml", styles)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    write_json()
    write_csv_and_txt()
    write_pdf()
    write_docx()
    write_xlsx()
    print(f"Generated full-invoice example inputs in {OUTPUT}")


if __name__ == "__main__":
    main()
